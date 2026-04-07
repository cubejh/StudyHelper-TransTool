import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

OUTPUT_PATH = "testingFolder/output.docx"

CHINESE_FONT = "標楷體"
ENGLISH_FONT = "Times New Roman"
ENGLISH_SIZE = Pt(12)
NO_ITALIC_WORDS = {"sin", "cos", "tan", "log", "exp", "sqrt"}

def contains_chinese(char):
    return '\u4e00' <= char <= '\u9fff'

def text_to_word(text: str, output_path: str = OUTPUT_PATH):
    doc = Document()

    for line in text.split("\n"):
        p = doc.add_paragraph()
        tokens = re.findall(r'\([A-Z]\)|\w+|[^\w\s]', line)  # split into words, symbols, or (A)

        for token in tokens:
            run = p.add_run(token)

            if any(contains_chinese(c) for c in token):
                run.font.name = CHINESE_FONT
                run._element.rPr.rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
                run.font.italic = False
            else:
                run.font.name = ENGLISH_FONT
                run._element.rPr.rFonts.set(qn("w:eastAsia"), ENGLISH_FONT)
                run.font.italic = False

                if re.fullmatch(r'\([A-Z]\)', token):
                    run.font.italic = False
                elif re.fullmatch(r'[A-Za-z]+', token) and token.lower() not in NO_ITALIC_WORDS:
                    run.font.italic = True

            run.font.size = ENGLISH_SIZE

    doc.save(output_path)